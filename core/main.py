"""
main.py - 试卷抄袭检测系统后端主服务

这是基于 FastAPI 的 RESTful API 服务，提供以下核心功能：
1. 试卷结构分析（解析 DOCX 文档为题目块）
2. 题目标签提取（通过 LLM 分析题目类型、分数、知识点）
3. 题目相似度检测（NLP 模型或 LLM 两种模式）
4. 试卷格式校验（科目名称、代码、页码、总分等）
5. 报告生成（将检测结果标注到原始文档）

依赖模块：
- ocr_module: 本地 OCR 引擎，处理图片中的文本识别
- llm_service: LLM 服务调用，包括题目分割、标签提取、相似度计算
- schemas: Pydantic 数据模型定义

技术栈：
- FastAPI: 高性能异步 Web 框架
- sentence-transformers: 本地 NLP 相似度计算
- python-docx: DOCX 文档读写
- pypandoc: DOCX 转 Markdown
- LibreOffice: 文档格式转换（PDF/DOCX）
"""

import asyncio
import difflib
import glob
import hashlib
import json
import math
import os
import re
import shutil
import subprocess
import tempfile
import sys
from contextlib import asynccontextmanager
from io import BytesIO
from typing import Optional, List, Dict, Any, Tuple, Union

import numpy as np
import torch
from sentence_transformers import SentenceTransformer  # 本地语义相似度模型
import pypandoc  # DOCX 转 Markdown 工具
from docx import Document  # DOCX 文档处理
from docx.oxml import OxmlElement  # DOCX XML 操作
from docx.shared import RGBColor  # DOCX 颜色设置
from docx.text.paragraph import Paragraph  # DOCX 段落对象
from fastapi import FastAPI, UploadFile, File, Form, HTTPException, Request  # Web 框架
from docx.oxml.ns import qn  # DOCX 命名空间

# 导入自定义模块
from llm_service import (
    get_online_segmentation,  # 在线 LLM 试卷分割
    extract_question_info,    # 题目信息提取
    check_tags_similarity,    # 标签相似度检测
    compute_question_similarity,  # 题目相似度计算
    call_llm_for_format_validation,  # 格式校验（上传文件）
    call_llm_for_format_validation_local  # 格式校验（本地解析）
)
from ocr_module import LocalOCREngine  # 本地 OCR 引擎
from schemas import (
    ExamMap, AppConfig, SimilarityRequest, SimilarityResponse,
    FormatValidationResponse, FormatCheckItemResult,
    ReportGenerateRequest
)


# ==================== 本地 NLP 相似度计算（基于 sentence-transformers） ====================
_NLP_MODEL: Optional[SentenceTransformer] = None
"""全局 NLP 模型实例，使用单例模式避免重复加载"""


def _get_nlp_model() -> SentenceTransformer:
    """
    单例模式加载多语言句子相似度模型（本地优先，否则在线下载）
    
    模型选择策略：
    1. 优先从本地路径 `./distiluse-base-multilingual-cased-v1` 加载
    2. 支持 HuggingFace Transformers 格式的快照目录结构
    3. 本地加载失败时，自动从 HuggingFace Hub 下载
    
    使用的模型：distiluse-base-multilingual-cased-v1
    - 支持多种语言（包括中文）
    - 模型较小，适合本地部署
    - 输出 512 维向量用于相似度计算
    
    返回:
        SentenceTransformer: 已加载的语义相似度模型
    """
    global _NLP_MODEL
    if _NLP_MODEL is not None:
        return _NLP_MODEL

    # 模型配置
    base_dir = "./distiluse-base-multilingual-cased-v1"  # 本地模型目录
    official_name = "sentence-transformers/distiluse-base-multilingual-cased-v1"  # HuggingFace 模型名
    device = 'cuda' if torch.cuda.is_available() else 'cpu'  # 自动检测计算设备

    # 自动寻找本地模型路径
    local_path = None
    if os.path.isdir(base_dir):
        # 检查是否为标准模型目录（包含 config.json）
        if os.path.exists(os.path.join(base_dir, "config.json")):
            local_path = base_dir
        else:
            # 检查 snapshots 子目录（Transformers 缓存格式）
            snapshots_dir = os.path.join(base_dir, "snapshots")
            if os.path.isdir(snapshots_dir):
                snapshots = [d for d in os.listdir(snapshots_dir)
                             if os.path.isdir(os.path.join(snapshots_dir, d))]
                if snapshots:
                    local_path = os.path.join(snapshots_dir, snapshots[0])

    # 尝试从本地加载
    if local_path:
        print(f"[NLP] 从本地加载模型: {local_path}")
        try:
            _NLP_MODEL = SentenceTransformer(local_path, device=device)
            print("[NLP] 本地模型加载成功")
            return _NLP_MODEL
        except Exception as e:
            print(f"[NLP] 本地加载失败，将尝试在线下载: {e}")

    # 从 HuggingFace 在线下载
    print("[NLP] 从 HuggingFace 在线下载模型...")
    _NLP_MODEL = SentenceTransformer(official_name, device=device)
    print("[NLP] 在线模型加载成功")
    return _NLP_MODEL


def _clean_text(text: str) -> str:
    """
    文本预处理函数，用于相似度比较前的文本清洗
    
    处理步骤：
    1. 确保输入为字符串类型
    2. 移除 HTML 标签
    3. 将多个连续空白字符替换为单个空格
    4. 去除首尾空白
    
    参数:
        text: 原始文本字符串
        
    返回:
        str: 清洗后的文本
    """
    if not isinstance(text, str):
        text = str(text)
    # 移除 HTML 标签（如果存在）
    cleaned = re.sub(r'<[^>]+>', '', text)
    # 合并连续空白字符
    cleaned = re.sub(r'\s+', ' ', cleaned).strip()
    return cleaned


def compute_local_similarity(text1: str, text2: str) -> float:
    """
    使用本地多语言 sentence-transformers 模型计算两个文本的余弦相似度
    
    计算流程：
    1. 加载预训练的语义模型
    2. 对两段文本进行预处理（清洗）
    3. 将文本编码为 512 维向量
    4. 计算向量之间的余弦相似度
    
    参数:
        text1: 第一段文本
        text2: 第二段文本
        
    返回:
        float: 余弦相似度值（范围 0~1），值越大表示越相似
    """
    model = _get_nlp_model()
    clean1 = _clean_text(text1)
    clean2 = _clean_text(text2)
    
    # 将文本编码为向量
    emb1 = model.encode([clean1], convert_to_numpy=True)[0]
    emb2 = model.encode([clean2], convert_to_numpy=True)[0]
    
    # 计算余弦相似度
    dot = np.dot(emb1, emb2)
    norm1 = np.linalg.norm(emb1)
    norm2 = np.linalg.norm(emb2)
    
    # 处理零向量情况
    if norm1 == 0 or norm2 == 0:
        return 0.0
    
    return float(dot / (norm1 * norm2))


# ==================== LibreOffice 便携版专用路径 ====================
def get_soffice_path() -> str:
    """
    获取 LibreOffice 可执行文件路径（仅使用同目录下的便携版）
    
    设计目的：实现完全便携化部署，不依赖系统安装的 LibreOffice。
    支持的目录结构：
    - Windows: LibreOfficePortable/App/libreoffice/program/soffice.exe
    - Linux: LibreOfficePortable/App/libreoffice/program/soffice
    
    返回:
        str: LibreOffice 可执行文件的完整路径
        
    抛出:
        RuntimeError: 未找到有效的 LibreOffice 便携版
    """
    script_dir = os.path.dirname(os.path.abspath(__file__))
    portable_root = os.path.join(script_dir, "LibreOfficePortable")

    # 根据操作系统选择候选路径
    if sys.platform == "win32":
        candidates = [
            os.path.join(portable_root, "App", "libreoffice", "program", "soffice.exe"),
            os.path.join(portable_root, "LibreOfficePortable.exe"),
        ]
    else:
        candidates = [
            os.path.join(portable_root, "App", "libreoffice", "program", "soffice"),
            os.path.join(portable_root, "libreoffice", "program", "soffice"),
        ]

    # 遍历候选路径，返回第一个存在的
    for cand in candidates:
        if os.path.isfile(cand):
            return cand

    # 未找到时抛出详细错误信息
    raise RuntimeError(
        f"未找到 LibreOffice 便携版，请确保目录存在且结构正确：{portable_root}\n"
        "期望路径例如：LibreOfficePortable/App/libreoffice/program/soffice.exe"
    )


_SOFFICE_PATH: Optional[str] = None
"""缓存的 LibreOffice 路径，避免重复查找"""


def get_soffice_path_cached() -> str:
    """
    获取缓存的 LibreOffice 路径（懒加载模式）
    
    使用单例模式缓存路径，避免每次调用都进行文件系统检查。
    
    返回:
        str: LibreOffice 可执行文件的完整路径
    """
    global _SOFFICE_PATH
    if _SOFFICE_PATH is None:
        _SOFFICE_PATH = get_soffice_path()
    return _SOFFICE_PATH


# ==================== 应用生命周期管理 ====================
@asynccontextmanager
async def lifespan(app: FastAPI):
    """
    FastAPI 应用生命周期管理器
    
    启动阶段（yield 之前）：
    1. 初始化本地 OCR 引擎（支持图片中的文本识别）
    2. 预加载本地 NLP 语义模型（避免首次请求延迟）
    3. 初始化应用配置对象
    
    关闭阶段（yield 之后）：
    1. 卸载 OCR 引擎，释放显存/内存资源
    
    参数:
        app: FastAPI 应用实例
    """
    global ocr_engine

    # 初始化 OCR 引擎（可能失败，失败时设置为 None）
    try:
        ocr_engine = LocalOCREngine()
        print("[OCR] 本地 OCR 引擎初始化成功")
    except Exception as e:
        ocr_engine = None
        print(f"[OCR] 本地 OCR 引擎初始化失败，将跳过图片识别: {e}")

    # 预加载本地 NLP 模型（可选，避免首次请求延迟）
    try:
        _get_nlp_model()
        print("[NLP] 本地 NLP 模型预加载成功")
    except Exception as e:
        print(f"[NLP] 本地 NLP 模型预加载失败: {e}")

    # 初始化应用状态
    app.state.config = AppConfig()  # 全局配置（API 地址、密钥等）
    app.state.last_exam_map = None  # 最近一次试卷分析结果
    app.state.last_blocks = None    # 最近一次文档切分结果
    
    yield  # 应用运行中

    # 应用关闭时的清理工作
    if ocr_engine:
        try:
            ocr_engine.unload()
            print("[OCR] OCR 引擎已卸载")
        except Exception as e:
            print(f"[OCR] OCR 引擎卸载失败: {e}")


# 创建 FastAPI 应用实例
app = FastAPI(
    title="试卷抄袭检测系统",
    description="基于 LLM 和 NLP 的试卷分析与抄袭检测服务",
    version="1.0.0",
    lifespan=lifespan
)

# 全局 OCR 引擎实例
ocr_engine: Optional[LocalOCREngine] = None


# ==================== 异步 OCR 替换函数 ====================
async def convert_image_blob_to_png(blob: bytes, cache_dir: Optional[str] = None) -> bytes:
    """
    将图片二进制数据转换为 PNG 格式（用于 OCR 识别）
    
    支持的输入格式：WMF、EMF、JPG、PNG、GIF
    使用 LibreOffice 进行格式转换，支持矢量图转位图
    
    特性：
    - 支持缓存机制，避免重复转换相同图片
    - 使用 MD5 哈希值作为缓存键
    
    参数:
        blob: 图片原始二进制数据
        cache_dir: 缓存目录路径（可选，默认使用系统临时目录）
        
    返回:
        bytes: PNG 格式的图片二进制数据
        
    抛出:
        RuntimeError: LibreOffice 转换失败
        FileNotFoundError: 转换后未找到 PNG 文件
    """
    soffice_path = get_soffice_path_cached()

    # 设置缓存目录
    if cache_dir is None:
        cache_dir = os.path.join(tempfile.gettempdir(), "ocr_img_cache")
    os.makedirs(cache_dir, exist_ok=True)

    # 计算图片哈希值用于缓存
    blob_hash = hashlib.md5(blob).hexdigest()
    cached_png_path = os.path.join(cache_dir, f"{blob_hash}.png")
    
    # 检查缓存是否存在
    if os.path.exists(cached_png_path):
        with open(cached_png_path, "rb") as f:
            return f.read()

    # 根据文件头判断图片格式
    if blob.startswith(b'\xd7\xcd\xc6\x9a'):
        ext = ".wmf"
    elif blob.startswith(b'\x01\x00\x00\x00'):
        ext = ".emf"
    elif blob.startswith(b'\xff\xd8'):
        ext = ".jpg"
    elif blob.startswith(b'\x89PNG'):
        ext = ".png"
    elif blob.startswith(b'GIF8'):
        ext = ".gif"
    else:
        ext = ".bin"

    # 使用临时目录进行转换
    with tempfile.TemporaryDirectory() as tmpdir:
        tmp_in_path = os.path.join(tmpdir, f"input{ext}")
        with open(tmp_in_path, "wb") as f:
            f.write(blob)

        out_dir = os.path.join(tmpdir, "output")
        os.makedirs(out_dir, exist_ok=True)

        # 构建 LibreOffice 转换命令
        cmd = [
            soffice_path, "--headless", "--convert-to", "png",
            "--outdir", out_dir, tmp_in_path
        ]

        # 异步执行转换命令
        proc = await asyncio.create_subprocess_exec(
            *cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE
        )
        stdout, stderr = await asyncio.wait_for(proc.communicate(), timeout=30)
        
        # 检查转换是否成功
        if proc.returncode != 0:
            error_msg = stderr.decode() if stderr else "Unknown error"
            raise RuntimeError(f"LibreOffice 转换失败 (code {proc.returncode}): {error_msg}")

        # 查找转换后的 PNG 文件
        base_name = os.path.splitext(os.path.basename(tmp_in_path))[0]
        png_path = os.path.join(out_dir, f"{base_name}.png")
        if not os.path.exists(png_path):
            files = glob.glob(os.path.join(out_dir, "*.png"))
            if not files:
                raise FileNotFoundError(f"未找到转换后的 PNG 文件，输出目录: {out_dir}")
            png_path = files[0]

        # 读取并缓存结果
        with open(png_path, "rb") as f:
            png_bytes = f.read()
        with open(cached_png_path, "wb") as f:
            f.write(png_bytes)
        
        return png_bytes


async def replace_images_with_ocr_async(docx_bytes: bytes, markdown_text: str, ocr_engine: LocalOCREngine) -> str:
    """
    异步将 Markdown 中的图片占位符替换为 OCR 识别结果
    
    处理流程：
    1. 从 DOCX 文件中提取所有图片数据
    2. 在 Markdown 文本中查找图片标记（![...](...)）
    3. 对每张图片进行 OCR 识别
    4. 将识别结果替换到原始位置
    
    参数:
        docx_bytes: DOCX 文件的二进制数据
        markdown_text: 从 DOCX 转换得到的 Markdown 文本
        ocr_engine: 已初始化的 OCR 引擎实例
        
    返回:
        str: 替换了图片内容的 Markdown 文本
    """
    # 从 DOCX 中提取所有图片
    doc = Document(BytesIO(docx_bytes))
    image_map: Dict[str, bytes] = {}
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            image_map[rel.target_ref] = rel.target_part.blob

    # 查找 Markdown 中的图片标记
    pattern = re.compile(r'!\[.*?\]\((.*?)\)')
    matches = list(pattern.finditer(markdown_text))
    if not matches:
        return markdown_text

    # 内部函数：识别单张图片
    async def recognize_image(blob: bytes, img_path: str) -> str:
        try:
            # 转换为 PNG 格式（OCR 引擎要求）
            png_bytes = await convert_image_blob_to_png(blob)
            try:
                text = await ocr_engine.recognize(png_bytes)
            except Exception:
                return ""
            return text.strip()
        except Exception:
            return ""

    # 并行处理所有图片识别任务
    tasks = []
    for match in matches:
        img_path = match.group(1)
        blob = image_map.get(img_path)
        
        # 如果直接匹配失败，尝试通过文件名匹配
        if blob is None:
            base_name = os.path.basename(img_path)
            for ref, b in image_map.items():
                if ref.endswith(base_name) or base_name in ref:
                    blob = b
                    break
        
        # 加入任务队列
        if blob is None:
            tasks.append(asyncio.sleep(0, result=""))
        else:
            tasks.append(recognize_image(blob, img_path))

    # 等待所有任务完成
    ocr_results = await asyncio.gather(*tasks)

    # 从后往前替换，避免索引偏移
    result_chars = list(markdown_text)
    for (match, ocr_text) in zip(reversed(matches), reversed(ocr_results)):
        start, end = match.start(), match.end()
        if ocr_text:
            replacement = f" {ocr_text} "
        else:
            replacement = match.group(0)  # 保持原图片标记
        result_chars[start:end] = replacement
    
    new_markdown = ''.join(result_chars)
    return new_markdown


# ==================== 文档转换辅助函数 ====================
def convert_docx_to_markdown(docx_bytes: bytes) -> str:
    """
    将 DOCX 文件转换为 Markdown 格式
    
    使用 pypandoc 进行格式转换，保留数学公式和原始 TeX 内容。
    
    参数:
        docx_bytes: DOCX 文件的二进制数据
        
    返回:
        str: Markdown 格式的文本内容
    """
    with tempfile.TemporaryDirectory() as tmpdir:
        tmp_in_path = os.path.join(tmpdir, "input.docx")
        with open(tmp_in_path, "wb") as f:
            f.write(docx_bytes)

        # 使用 pypandoc 转换，保留数学公式
        output = pypandoc.convert_file(
            tmp_in_path,
            'markdown+tex_math_dollars+raw_tex',
            format='docx',
            extra_args=['--wrap=none']
        )
        return output


def convert_docx_to_pdf(docx_bytes: bytes) -> bytes:
    """
    将 DOCX 文件转换为 PDF 格式（用于 LLM 格式校验）
    
    使用 LibreOffice 进行转换，支持复杂格式和数学公式。
    
    参数:
        docx_bytes: DOCX 文件的二进制数据
        
    返回:
        bytes: PDF 文件的二进制数据
        
    抛出:
        RuntimeError: 转换失败或未找到输出文件
    """
    soffice_path = get_soffice_path_cached()
    
    with tempfile.TemporaryDirectory() as tmpdir:
        docx_path = os.path.join(tmpdir, "input.docx")
        with open(docx_path, "wb") as f:
            f.write(docx_bytes)

        # 构建转换命令
        cmd = [
            soffice_path, "--headless", "--convert-to", "pdf",
            "--outdir", tmpdir, docx_path
        ]
        result = subprocess.run(cmd, capture_output=True, text=True)
        
        # 检查转换结果
        if result.returncode != 0:
            raise RuntimeError(f"LibreOffice 转换失败: {result.stderr}")

        # 读取转换后的 PDF
        pdf_path = os.path.join(tmpdir, "input.pdf")
        if not os.path.exists(pdf_path):
            raise RuntimeError("转换后未找到 PDF 文件")

        with open(pdf_path, "rb") as f:
            pdf_bytes = f.read()
        
        return pdf_bytes


def split_markdown_into_blocks(markdown_text: str) -> List[Dict[str, Any]]:
    """
    将 Markdown 文本切分为语义块
    
    按空行分隔文本，每个块标记为普通文本或数学公式。
    
    参数:
        markdown_text: Markdown 格式的文本
        
    返回:
        List[Dict]: 块列表，每个块包含 id、text、tag、image、mime 字段
    """
    lines = markdown_text.split('\n')
    blocks: List[Dict[str, Any]] = []
    current_lines: List[str] = []
    block_id = 0

    # 内部函数：将当前累积的行输出为一个块
    def flush_block() -> None:
        nonlocal block_id
        if not current_lines:
            return
        text = '\n'.join(current_lines).strip()
        if text:
            # 判断块类型：单行 $$...$$ 格式为数学公式
            tag = 'math' if text.startswith('$$') and text.endswith('$$') and len(text.split('\n')) == 1 else 'normal'
            blocks.append({
                "id": block_id,
                "text": text,
                "tag": tag,
                "image": None,
                "mime": None
            })
            block_id += 1
        current_lines.clear()

    # 遍历所有行，按空行分割
    for line in lines:
        if line.strip() == '':
            flush_block()
        else:
            current_lines.append(line)
    flush_block()
    
    return blocks


def get_blocks_from_docx(docx_path: str) -> List[Dict[str, Any]]:
    """
    从 DOCX 文件路径读取并切分为语义块
    
    参数:
        docx_path: DOCX 文件的完整路径
        
    返回:
        List[Dict]: 块列表
    """
    with open(docx_path, "rb") as f:
        docx_bytes = f.read()
    markdown_text = convert_docx_to_markdown(docx_bytes)
    blocks = split_markdown_into_blocks(markdown_text)
    return blocks


# ==================== API 端点 ====================
@app.post("/set-config")
async def set_config(
        request: Request,
        base_url: str = Form(...),
        model_name: str = Form(...),
        api_key: str = Form(...)
) -> Dict[str, Any]:
    """
    设置全局 API 配置
    
    配置 LLM 服务的连接信息，后续的 LLM 相关操作将使用此配置。
    
    请求参数（表单）:
        base_url: LLM API 服务地址（如 http://localhost:11434/v1）
        model_name: 模型名称（如 llama3.1:8b、gpt-4o 或 nlp）
        api_key: API 密钥（本地模型可设为任意值）
    
    返回:
        {"status": "success", "data": {...}} - 配置信息
    """
    config: AppConfig = request.app.state.config
    config.base_url = base_url.strip()
    config.model_name = model_name.strip()
    config.api_key = api_key.strip()
    return {"status": "success", "data": config.model_dump()}


@app.post("/analyze-exam")
async def analyze_exam(
        request: Request,
        file: UploadFile = File(...),
        api_type: str = Form(default="online"),
        prompt: str = Form(default="")
) -> Dict[str, Any]:
    """
    分析试卷结构，将文档切分为题目
    
    处理流程：
    1. 读取 DOCX 文件并转换为 Markdown
    2. 对图片内容进行 OCR 识别（如果 OCR 引擎可用）
    3. 将文本切分为语义块
    4. 使用 LLM 分析块结构，识别题目边界
    
    请求参数（表单）:
        file: DOCX 文件
        api_type: 分析模式（online/llm/none）
        prompt: LLM 提示词，定义试卷分割规则
    
    返回:
        {
            "status": "success",
            "data": {
                "api_type": str,
                "blocks": List[Dict],  // 切分后的语义块
                "questions": List[QuestionGroup]  // 识别出的题目列表
            }
        }
    
    错误:
        400 - 文件格式错误、配置未完成、缺少 prompt
        500 - 服务器内部错误
    """
    # 验证文件格式
    if not file.filename.endswith('.docx'):
        raise HTTPException(400, "仅支持 .docx 文件")
    
    # 验证配置
    config = request.app.state.config
    if api_type == "online" and not config.is_online_ready():
        raise HTTPException(400, "在线模式未配置：请先调用 /set-config")
    if not prompt.strip():
        raise HTTPException(400, "必须提供 prompt")

    # 验证模型选择
    if api_type == "online" and config.model_name.strip().lower() == "nlp":
        raise HTTPException(400, "NLP模型无法处理试卷分割问题")

    try:
        # 读取并转换文档
        docx_bytes = await file.read()
        markdown_text = convert_docx_to_markdown(docx_bytes)

        # OCR 图片内容（如果引擎可用）
        if ocr_engine is not None:
            markdown_text = await replace_images_with_ocr_async(docx_bytes, markdown_text, ocr_engine)

        # 切分为语义块
        blocks = split_markdown_into_blocks(markdown_text)

        # 分析题目结构
        if api_type == "none":
            exam_map = ExamMap(questions=[])
        else:
            exam_map = get_online_segmentation(
                blocks=blocks,
                api_key=config.api_key,
                prompt=prompt,
                base_url=config.base_url,
                model_name=config.model_name
            )

        # 保存到应用状态供后续使用
        request.app.state.last_exam_map = exam_map
        request.app.state.last_blocks = blocks
        
        # 为每个题目生成标题（从第一个块提取）
        block_map = {b["id"]: b for b in blocks}
        for q in exam_map.questions:
            if q.indices:
                first_block = block_map.get(q.indices[0])
                if first_block:
                    raw_text = first_block.get("text", "")
                    match = re.search(r'^(\d+)[\.\．]\s*(.+?)$', raw_text, re.MULTILINE)
                    q.title = (match.group(2).strip()[:50] if match else raw_text[:50]) or "题目"
                else:
                    q.title = "题目"

        # 清理块数据（移除图片字段）
        clean_blocks = [{k: v for k, v in b.items() if k != "image"} for b in blocks]

        return {"status": "success",
                "data": {"api_type": api_type, "blocks": clean_blocks, "questions": exam_map.questions}}
    except Exception as e:
        raise HTTPException(500, f"服务器内部错误: {str(e)}")


@app.post("/extract-tags")
async def extract_tags_endpoint(
        request: Request,
        content: str = Form(...),
        context_blocks: Optional[str] = Form(None),
        question_id: Optional[int] = Form(None)
) -> Dict[str, Any]:
    """
    提取题目信息（题型、分数、知识点标签）
    
    使用 LLM 分析题目内容，提取：
    - 题型（选择题、填空题、计算题等）
    - 分数（从题目或上下文提取）
    - 三个知识点标签
    
    请求参数（表单）:
        content: 题目内容文本
        context_blocks: 上下文文本列表（JSON 字符串，用于分数推断）
        question_id: 题目ID（用于获取默认分数）
    
    返回:
        {
            "status": "success",
            "data": {
                "题型": str,      // 题目类型
                "分数": float,    // 题目分值
                "标签1": str,     // 知识点标签1
                "标签2": str,     // 知识点标签2
                "标签3": str      // 知识点标签3
            }
        }
    
    错误:
        400 - 配置未完成、内容为空、模型不支持
        500 - 服务器内部错误
    """
    config: AppConfig = request.app.state.config
    
    # 验证配置
    if not config.is_online_ready():
        raise HTTPException(400, "全局配置未完成：请先调用 /set-config")
    if not content.strip():
        raise HTTPException(400, "题目内容不能为空")

    # 验证模型选择
    if config.model_name.strip().lower() == "nlp":
        raise HTTPException(400, "NLP模型无法处理标签提取问题")

    # 解析上下文文本（用于分数推断）
    context_texts: Optional[List[str]] = None
    if context_blocks and context_blocks.strip():
        try:
            context_texts = json.loads(context_blocks)
            if not isinstance(context_texts, list):
                context_texts = None
        except Exception:
            pass

    # 获取默认分数（从上次分析结果中）
    default_score: Optional[float] = None
    if question_id is not None and hasattr(request.app.state, 'last_exam_map'):
        exam_map = request.app.state.last_exam_map
        if exam_map:
            for q in exam_map.questions:
                if q.id == question_id:
                    default_score = getattr(q, 'default_score', None)
                    break

    try:
        # 调用 LLM 提取题目信息
        info = extract_question_info(
            content=content,
            api_key=config.api_key,
            base_url=config.base_url,
            model_name=config.model_name,
            context_blocks=context_texts
        )
        
        # 如果 LLM 返回分数为 0，使用默认分数
        if info.score == 0.0 and default_score is not None:
            info.score = float(default_score)
        
        return {"status": "success", "data": info.model_dump(by_alias=True)}
    except Exception as e:
        raise HTTPException(500, f"服务器内部错误: {str(e)}")


@app.post("/check-similarity", response_model=SimilarityResponse)
async def check_similarity(request: Request, payload: SimilarityRequest) -> SimilarityResponse:
    """"""
    检查两道题目的相似度（重复检测）
    
    支持两种模式：
    1. NLP 模式：使用本地 sentence-transformers 模型计算语义相似度
    2. LLM 模式：使用 LLM 进行深度语义分析，包含标签检查
    
    请求体（JSON）:
        {
            "question1": {
                "body": str,      // 题目正文
                "type": str,      // 题型
                "tags": List[str] // 标签列表
            },
            "question2": {...},   // 第二道题目
            "threshold": float    // 重复阈值（0-100）
        }
    
    返回:
        SimilarityResponse:
            is_duplicate: bool    // 是否判定为重复
            similarity: float     // 相似度百分比（0-100）
            reason: Optional[str] // 重复原因（相似度>=阈值时）
            tag_check_passed: bool // 标签检查是否通过
    
    错误:
        400 - 配置未完成
        500 - 服务器内部错误
    """
    try:
        config: AppConfig = request.app.state.config
        
        # 验证配置
        if not config.is_online_ready():
            raise HTTPException(400, "全局配置未完成：请先调用 /set-config")

        q1, q2 = payload.question1, payload.question2

        # ===== NLP 模式 =====
        # 使用本地 sentence-transformers 模型进行快速相似度计算
        if config.model_name.strip().lower() == "nlp":
            try:
                # 计算余弦相似度并应用非线性映射（提高区分度）
                sim_score_decimal = compute_local_similarity(q1.body, q2.body)
                sim_score_percent = sim_score_decimal * 100.0
                sim_score_percent = 100 * (sim_score_percent / 100) ** 5.5
            except Exception as e:
                raise HTTPException(500, f"本地 NLP 模型计算相似度失败: {type(e).__name__}: {str(e)}")

            is_dup = sim_score_percent >= payload.threshold
            print(f"[NLP] 阈值: {payload.threshold}%, 相似度: {sim_score_percent:.2f}%, 判定: {'重复' if is_dup else '不重复'}")

            reason = None
            if is_dup:
                reason = f"基于 NLP 模型的内容相似度判定（相似度 {sim_score_percent:.2f}%）"
            
            return SimilarityResponse(
                is_duplicate=is_dup,
                similarity=sim_score_percent,
                reason=reason,
                tag_check_passed=True
            )

        # ===== LLM 模式 =====
        # 先进行标签相似度检查（快速过滤）
        tag_pass = check_tags_similarity(
            tags1=q1.tags, tags2=q2.tags,
            api_key=config.api_key, base_url=config.base_url, model_name=config.model_name
        )
        if not tag_pass:
            return SimilarityResponse(is_duplicate=False, similarity=0.0, reason=None, tag_check_passed=False)

        # 使用 LLM 进行深度相似度分析
        similarity, reason = compute_question_similarity(
            q1, q2,
            api_key=config.api_key,
            base_url=config.base_url,
            model_name=config.model_name,
            reason_threshold=payload.threshold
        )
        
        is_dup = similarity >= payload.threshold
        return SimilarityResponse(
            is_duplicate=is_dup,
            similarity=float(similarity),
            reason=reason if is_dup else None,
            tag_check_passed=True
        )

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(500, f"服务器内部错误: {type(e).__name__}: {str(e)}")

@app.post("/validate-format", response_model=FormatValidationResponse)
async def validate_format(
        request: Request,
        file: UploadFile = File(...),
        format_checks: str = Form(...),
        expected_total_score: Optional[float] = Form(None),
        is_local_api: bool = Form(False)
) -> FormatValidationResponse:
    """
    校验试卷格式规范
    
    检查试卷的格式项，包括：
    - 科目名称是否正确填写
    - 科目代码是否正确填写
    - 科目名称与代码是否匹配
    - 页眉科目代码是否一致
    - 页码格式是否正确
    - 章节是否均标记分值
    - 序号是否正确填写
    - 选择题选项是否重复
    - 总分是否在预设范围内
    
    请求参数（表单）:
        file: DOCX 文件
        format_checks: 启用的校验项（JSON 字符串，如 {"SUBJECT_NAME_CORRECT_FILLED": true, ...}）
        expected_total_score: 预期总分（用于总分校验）
        is_local_api: 是否使用本地解析模式（不上传文件）
    
    返回:
        FormatValidationResponse:
            results: List[FormatCheckItemResult] - 各项校验结果
            summary: str - 总体结论
    
    错误:
        400 - 文件格式错误、参数无效、配置未完成、模型不支持
        502 - 无法连接到 LLM 服务
        500 - 服务器内部错误
    """
    # 验证文件格式
    if not file.filename.endswith('.docx'):
        raise HTTPException(400, "仅支持 .docx 文件")
    
    # 解析校验项配置
    try:
        checks_dict = json.loads(format_checks)
    except json.JSONDecodeError:
        raise HTTPException(400, "format_checks 参数必须是有效的 JSON 字符串")

    # 读取文件
    docx_bytes = await file.read()
    original_filename = file.filename
    
    # 验证配置
    config: AppConfig = request.app.state.config
    if not config.is_online_ready():
        raise HTTPException(400, "未配置 LLM 服务，请先调用 /set-config")

    # 验证模型选择
    if config.model_name.strip().lower() == "nlp":
        raise HTTPException(400, "NLP模型无法处理格式校验问题")

    try:
        # 根据模式选择校验方式
        if is_local_api:
            # 本地解析模式：先转换为文本再发送给 LLM
            results = call_llm_for_format_validation_local(
                docx_bytes=docx_bytes,
                enabled_checks=checks_dict,
                api_key=config.api_key,
                base_url=config.base_url,
                model_name=config.model_name,
                expected_total_score=expected_total_score
            )
        else:
            # 上传文件模式：转换为 PDF 后上传到 LLM 服务
            pdf_bytes = convert_docx_to_pdf(docx_bytes)
            pdf_filename = original_filename.rsplit('.', 1)[0] + ".pdf"
            results = call_llm_for_format_validation(
                file_bytes=pdf_bytes,
                filename=pdf_filename,
                enabled_checks=checks_dict,
                api_key=config.api_key,
                base_url=config.base_url,
                model_name=config.model_name,
                expected_total_score=expected_total_score
            )
    except HTTPException:
        raise
    except Exception as e:
        error_msg = str(e)
        # 处理连接错误
        if "Connection error" in error_msg or "Connection refused" in error_msg or "timed out" in error_msg:
            raise HTTPException(
                status_code=502,
                detail=f"无法连接到 LLM 服务 (base_url: {config.base_url})。请检查：\n"
                       "1. 本地 LLM 服务是否已启动（如 Ollama、LM Studio）\n"
                       "2. base_url 是否正确（例如 Ollama 默认 http://localhost:11434/v1）\n"
                       "3. 端口是否开放，防火墙是否放行\n"
                       f"原始错误: {error_msg}"
            )
        raise HTTPException(500, f"格式校验失败: {error_msg}")

    # 统计结果并构建响应
    passed_count = sum(1 for r in results if r.get("passed", False))
    total_count = len(results)
    summary = f"通过 {passed_count} 项，未通过 {total_count - passed_count} 项"
    
    response_items = [
        FormatCheckItemResult(
            code=r.get("code", ""),
            name=r.get("name", ""),
            passed=r.get("passed", False),
            reason=r.get("reason", "")
        )
        for r in results
    ]
    
    return FormatValidationResponse(results=response_items, summary=summary)


@app.post("/generate-report")
async def generate_report(request: ReportGenerateRequest) -> Dict[str, Any]:
    """
    生成检测报告（将结果标注到原始 DOCX 文档）
    
    处理流程：
    1. 复制原始 DOCX 文件到目标路径
    2. 切分文档为语义块
    3. 在文档开头插入格式校验结果表格
    4. 对重复内容进行红色标红并添加注释
    
    请求体（JSON）:
        {
            "docx_file_path": str,           // 原始 DOCX 文件路径
            "format_checks": List[FormatCheckItemInput],  // 格式校验结果
            "duplicate_results": List[DuplicateResultInput],  // 重复检测结果
            "save_path": str                 // 报告保存路径（.docx）
        }
    
    返回:
        {
            "status": "success",
            "message": "报告生成成功",
            "save_path": str
        }
    
    错误:
        400 - 文件不存在、保存路径无效
        500 - 文件复制失败、文档处理失败、报告生成失败
    """
    # 验证文件路径
    if not os.path.exists(request.docx_file_path):
        raise HTTPException(400, f"原始文件不存在: {request.docx_file_path}")
    
    # 验证保存路径
    if not request.save_path.endswith('.docx'):
        raise HTTPException(400, "保存路径必须以 .docx 结尾")
    
    # 确保保存目录存在
    save_dir = os.path.dirname(request.save_path)
    if save_dir and not os.path.exists(save_dir):
        os.makedirs(save_dir)

    # 复制原始文件到目标路径
    try:
        shutil.copy2(request.docx_file_path, request.save_path)
    except Exception as e:
        raise HTTPException(500, f"复制文件失败: {str(e)}")

    # 切分文档为语义块（用于定位重复内容）
    try:
        blocks = get_blocks_from_docx(request.docx_file_path)
        block_text_map = {block["id"]: block["text"] for block in blocks}
    except Exception as e:
        if os.path.exists(request.save_path):
            os.remove(request.save_path)
        raise HTTPException(500, f"文档切分失败: {str(e)}")

    try:
        # 打开目标文档
        doc = Document(request.save_path)

        # WordprocessingML 命名空间
        w_ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

        # 获取文档中所有段落（包括内容控件内的段落）
        def get_all_paragraphs_with_location(doc: Document):
            items = []
            # 普通段落
            for para in doc.paragraphs:
                items.append((para.text, para, False))

            # 内容控件内的段落
            sdt_elements = doc.element.findall(f'.//{{{w_ns}}}sdt')
            for sdt in sdt_elements:
                para_elems = sdt.findall(f'.//{{{w_ns}}}p')
                for p_elem in para_elems:
                    texts = []
                    for t_elem in p_elem.findall(f'.//{{{w_ns}}}t'):
                        if t_elem.text:
                            texts.append(t_elem.text)
                    para_text = ''.join(texts)
                    items.append((para_text, p_elem, True))
            return items

        text_items = get_all_paragraphs_with_location(doc)

        # ========== 添加格式检查表格 + 分页符 ==========
        if request.format_checks:
            # 创建表格
            table = doc.add_table(rows=len(request.format_checks) + 1, cols=3)
            table.style = 'Table Grid'
            
            # 设置表头
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "格式条目"
            hdr_cells[1].text = "通过/不通过"
            hdr_cells[2].text = "原因"
            
            # 填充表格内容
            for i, check in enumerate(request.format_checks, start=1):
                row_cells = table.rows[i].cells
                row_cells[0].text = check.item
                row_cells[1].text = "✅ 通过" if check.passed else "❌ 不通过"
                row_cells[2].text = check.reason

            # 将表格移动到文档开头
            tbl = table._element
            body = doc.element.body
            first = body[0] if len(body) > 0 else None
            if first is not None:
                body.insert(0, tbl)
            else:
                body.append(tbl)

            # 在表格后面插入分页符
            page_break = OxmlElement('w:p')
            run = OxmlElement('w:r')
            br = OxmlElement('w:br')
            br.set(qn('w:type'), 'page')
            run.append(br)
            page_break.append(run)

            # 找到表格位置并插入分页符
            body = doc.element.body
            tbl_index = list(body).index(tbl)
            body.insert(tbl_index + 1, page_break)

        # ========== 标注重复内容 ==========
        for dup in reversed(request.duplicate_results):
            if not dup.block_ids:
                continue
            
            # 获取第一个块的文本（用于匹配）
            first_block_id = dup.block_ids[0]
            block_text = block_text_map.get(first_block_id)
            if not block_text:
                continue

            # 在文档中查找最匹配的段落
            best_item = None
            best_score = 0.0
            target_norm = re.sub(r'\s+', ' ', block_text.strip())
            
            for item_text, location, in_sdt in text_items:
                item_norm = re.sub(r'\s+', ' ', item_text.strip())
                score = difflib.SequenceMatcher(None, target_norm, item_norm).ratio()
                if score > best_score:
                    best_score = score
                    best_item = (item_text, location, in_sdt)

            if best_item is None:
                continue

            _, location, in_sdt = best_item

            # 标红段落文本（仅对普通段落）
            if not in_sdt and isinstance(location, Paragraph):
                for run in location.runs:
                    run.font.color.rgb = RGBColor(255, 0, 0)

            # 添加注释段落
            target_elem = location._element if isinstance(location, Paragraph) else location
            comment_text = f"【重复提示】相似度：{dup.similarity}% | 重复位置：{dup.duplicate_location} | 原因：{dup.reason}"

            new_para = OxmlElement('w:p')
            run = OxmlElement('w:r')
            t = OxmlElement('w:t')
            t.text = comment_text
            run.append(t)

            # 设置红色字体
            rPr = OxmlElement('w:rPr')
            color = OxmlElement('w:color')
            color.set(qn('w:val'), 'FF0000')
            rPr.append(color)
            run.insert(0, rPr)

            new_para.append(run)
            target_elem.addnext(new_para)

        # 保存文档
        doc.save(request.save_path)

        return {
            "status": "success",
            "message": "报告生成成功",
            "save_path": request.save_path
        }

    except Exception as e:
        import traceback
        traceback.print_exc()
        # 清理临时文件
        if os.path.exists(request.save_path):
            os.remove(request.save_path)
        raise HTTPException(500, f"报告生成失败: {str(e)}")


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="127.0.0.1", port=8000)